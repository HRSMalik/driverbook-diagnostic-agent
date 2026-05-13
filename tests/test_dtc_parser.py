"""
DTC Parser — unit test suite.

Tests _extract_fmi and parse_dtc_records with fixed input/output pairs.
Writes results to tests/reports/report_dtc_parser.docx.

Run from diagnostic_agent/:
    conda run -n driverbook python tests/test_dtc_parser.py
"""

import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.dtc_parser import _extract_fmi, parse_dtc_records

DOC_PATH = os.path.join(os.path.dirname(__file__), "reports", "report_dtc_parser.docx")

_GREEN = RGBColor(0x00, 0xAA, 0x55)
_RED   = RGBColor(0xCC, 0x00, 0x00)

# ============================================================================
# Test cases — (inputs, expected)
# ============================================================================

FMI_CASES: list[tuple] = [
    ("Out of Calibration FMI 13",  13),
    ("fmi 0 detected in module",   0),
    ("Voltage High FMI 3",         3),
    ("No fault info here",         None),
    ("",                           None),
    ("FMI",                        None),
    ("fmi  9 extra space",         9),
]

PARSE_CASES: list[tuple] = [
    (
        {
            "mil": True,
            "dtcs": {
                "SPN 521133": {"ecu": "Engine #2",          "desc": "Out of Calibration FMI 13"},
                "spn 0":      {"ecu": "Communications Unit", "desc": "No description"},
            },
        },
        "TRUCK-001",
        [
            {"code": "SPN 521133", "ecu": "Engine #2",           "fmi": 13,   "mil": True},
            {"code": "SPN 0",      "ecu": "Communications Unit",  "fmi": None, "mil": True},
        ],
    ),
    (
        {"mil": False, "dtcs": {}},
        "TRUCK-002",
        [],
    ),
    (
        {"dtcs": {"P0128": {"ecu": "ECM", "desc": "Coolant Temp Below Thermostat"}}},
        "TRUCK-003",
        [{"code": "P0128", "ecu": "ECM", "fmi": None, "mil": False}],
    ),
    (
        {"dtcs": {"  spn 929 ": {"ecu": "  Engine  ", "desc": ""}}},
        "TRUCK-004",
        [{"code": "SPN 929", "ecu": "Engine", "fmi": None}],
    ),
]


# ============================================================================
# Runners — return list of (label, status, diffs)
# ============================================================================

def run_fmi_cases() -> list[tuple]:
    rows = []
    for desc, expected in FMI_CASES:
        result = _extract_fmi(desc)
        ok = result == expected
        label = repr(desc)[:50]
        diffs = [] if ok else [f"expected={expected!r}  actual={result!r}"]
        rows.append((label, "PASS" if ok else "FAIL", diffs,
                     {"Input": repr(desc), "Expected": str(expected), "Actual": str(result)}))
    return rows


def _check_parse(result: list[dict], expected_shapes: list[dict]) -> list[str]:
    if len(result) != len(expected_shapes):
        return [f"length: expected {len(expected_shapes)} got {len(result)}"]
    diffs = []
    result_by_code = {r["code"]: r for r in result}
    for shape in expected_shapes:
        code = shape["code"]
        actual = result_by_code.get(code)
        if actual is None:
            diffs.append(f"missing code {code!r}")
            continue
        for key, exp_val in shape.items():
            if actual.get(key) != exp_val:
                diffs.append(f"[{code}] {key}: expected={exp_val!r} actual={actual.get(key)!r}")
    return diffs


def run_parse_cases() -> list[tuple]:
    rows = []
    for dtc_records, vehicle_id, expected_shapes in PARSE_CASES:
        result = parse_dtc_records(dtc_records, vehicle_id, timestamp="2025-01-01T00:00:00+00:00")
        diffs = _check_parse(result, expected_shapes)
        ok = len(diffs) == 0
        label = f"vehicle={vehicle_id}  expected_faults={len(expected_shapes)}"
        rows.append((label, "PASS" if ok else "FAIL", diffs,
                     {"Vehicle": vehicle_id,
                      "Expected faults": str(len(expected_shapes)),
                      "Actual faults":   str(len(result))}))
    return rows


# ============================================================================
# Console print
# ============================================================================

def _print_rows(section: str, rows: list[tuple]) -> int:
    print(f"=== {section} ===")
    passed = 0
    for label, status, diffs, _ in rows:
        print(f"  [{status}] {label}")
        for d in diffs:
            print(f"         x {d}")
        if status == "PASS":
            passed += 1
    return passed


# ============================================================================
# Doc generation
# ============================================================================

def _add_section(doc: Document, title: str, rows: list[tuple], col_keys: list[str]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(col_keys) + 2)
    table.style = "Table Grid"
    headers = col_keys + ["Status", "Notes"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for label, status, diffs, fields in rows:
        row = table.add_row().cells
        for i, key in enumerate(col_keys):
            row[i].text = fields.get(key, "")
        status_cell = row[len(col_keys)]
        run = status_cell.paragraphs[0].add_run(status)
        run.bold = True
        run.font.color.rgb = _GREEN if status == "PASS" else _RED
        row[len(col_keys) + 1].text = "; ".join(diffs) if diffs else ""

    doc.add_paragraph()


def generate_doc(fmi_rows: list[tuple], parse_rows: list[tuple], passed: int, total: int) -> None:
    os.makedirs(os.path.dirname(DOC_PATH), exist_ok=True)
    doc = Document()

    title = doc.add_heading("DTC Parser — Test Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Module: core/dtc_parser.py")
    p = doc.add_paragraph()
    run = p.add_run(f"Result: {passed}/{total} passed")
    run.bold = True
    run.font.color.rgb = _GREEN if passed == total else _RED
    doc.add_paragraph()

    _add_section(doc, "_extract_fmi", fmi_rows, ["Input", "Expected", "Actual"])
    _add_section(doc, "parse_dtc_records", parse_rows, ["Vehicle", "Expected faults", "Actual faults"])

    doc.save(DOC_PATH)
    print(f"Report written to {DOC_PATH}")


# ============================================================================
# Main
# ============================================================================

def main() -> None:
    fmi_rows   = run_fmi_cases()
    parse_rows = run_parse_cases()

    fmi_pass   = _print_rows("_extract_fmi",      fmi_rows)
    parse_pass = _print_rows("parse_dtc_records",  parse_rows)

    total_pass = fmi_pass + parse_pass
    total      = len(fmi_rows) + len(parse_rows)
    print(f"\n{total_pass}/{total} passed")

    generate_doc(fmi_rows, parse_rows, total_pass, total)


if __name__ == "__main__":
    main()

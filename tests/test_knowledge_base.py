"""
Knowledge Base — unit test suite.

Tests lookup, increment_occurrence, extract_and_insert_from_document,
and auto_learn_from_diagnosis using mongomock (in-memory MongoDB).
Writes results to tests/reports/report_knowledge_base.docx.

Run from diagnostic_agent/:
    conda run -n driverbook python tests/test_knowledge_base.py
"""

import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import mongomock

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.knowledge_base import (
    lookup,
    increment_occurrence,
    extract_and_insert_from_document,
    auto_learn_from_diagnosis,
)

DOC_PATH = os.path.join(os.path.dirname(__file__), "reports", "report_knowledge_base.docx")

_GREEN = RGBColor(0x00, 0xAA, 0x55)
_RED   = RGBColor(0xCC, 0x00, 0x00)


# ============================================================================
# Helpers
# ============================================================================

def _fresh_db():
    client = mongomock.MongoClient()
    db = client["test_diagnostics"]
    db["knowledge_base"].create_index("code", unique=True)
    return db


def _seed(db, code: str, source: str = "seed", **kwargs) -> None:
    doc = {"code": code, "source": source, "occurrence_count": 0, **kwargs}
    db["knowledge_base"].insert_one(doc)


# ============================================================================
# Test runners — each returns list of (label, status, diffs, fields)
# ============================================================================

def test_lookup() -> list[tuple]:
    rows = []

    def _case(label: str, condition: bool, detail: str = "", **fields) -> None:
        ok = bool(condition)
        rows.append((label, "PASS" if ok else "FAIL",
                     [] if ok else [detail], {"Test": label, "Notes": detail or ""}))

    db = _fresh_db()
    _seed(db, "SPN 521133", meaning="Calibration error")
    entry = lookup(db, "SPN 521133")
    _case("hit exact", entry is not None and entry.get("meaning") == "Calibration error")

    db = _fresh_db()
    _seed(db, "SPN 521133")
    _case("hit case-insensitive", lookup(db, "spn 521133") is not None)

    db = _fresh_db()
    _seed(db, "SPN 521133")
    _case("hit whitespace stripped", lookup(db, "  SPN 521133  ") is not None)

    db = _fresh_db()
    _case("miss returns None", lookup(db, "SPN 9999") is None)

    db = _fresh_db()
    _seed(db, "SPN 0")
    entry = lookup(db, "SPN 0")
    _case("_id excluded from result", entry is not None and "_id" not in entry,
          "_id was present in returned doc")

    return rows


def test_increment_occurrence() -> list[tuple]:
    rows = []

    def _case(label: str, condition: bool, detail: str = "") -> None:
        ok = bool(condition)
        rows.append((label, "PASS" if ok else "FAIL",
                     [] if ok else [detail], {"Test": label, "Notes": detail or ""}))

    db = _fresh_db()
    _seed(db, "SPN 100", occurrence_count=3)
    increment_occurrence(db, "SPN 100")
    doc = db["knowledge_base"].find_one({"code": "SPN 100"})
    _case("count incremented to 4", doc and doc.get("occurrence_count") == 4,
          f"got {doc.get('occurrence_count') if doc else None}")

    increment_occurrence(db, "SPN 100")
    doc = db["knowledge_base"].find_one({"code": "SPN 100"})
    _case("count incremented to 5 on second call", doc and doc.get("occurrence_count") == 5)

    db = _fresh_db()
    _seed(db, "SPN 200", occurrence_count=0)
    increment_occurrence(db, "spn 200")
    doc = db["knowledge_base"].find_one({"code": "SPN 200"})
    _case("increment case-insensitive", doc and doc.get("occurrence_count") == 1)

    return rows


def test_extract_and_insert() -> list[tuple]:
    rows = []
    fault = {"code": "SPN 9001", "ecu": "Engine #3", "fmi": 7, "description": "Voltage Low"}

    def _case(label: str, condition: bool, detail: str = "") -> None:
        ok = bool(condition)
        rows.append((label, "PASS" if ok else "FAIL",
                     [] if ok else [detail], {"Test": label, "Notes": detail or ""}))

    db = _fresh_db()
    inserted = extract_and_insert_from_document(db, fault)
    _case("first insert returns True", inserted is True)

    doc = db["knowledge_base"].find_one({"code": "SPN 9001"})
    _case("document created", doc is not None)
    _case("source=extracted_from_doc", doc and doc.get("source") == "extracted_from_doc")
    _case("meaning from description", doc and doc.get("meaning") == "Voltage Low")

    inserted_again = extract_and_insert_from_document(db, fault)
    _case("duplicate insert returns False", inserted_again is False)

    db = _fresh_db()
    _case("empty code returns False", extract_and_insert_from_document(db, {"code": ""}) is False)

    return rows


def test_auto_learn() -> list[tuple]:
    rows = []
    fault      = {"code": "SPN 8888", "ecu": "Transmission"}
    diagnostic = {
        "purpose": "Gear ratio error",
        "issue": "Slipping clutch pack",
        "severity": "High",
        "urgency": "Immediate Action",
        "explanation": "Clutch pack worn beyond spec",
        "resolution_steps": ["Inspect clutch", "Replace if worn"],
    }

    def _case(label: str, condition: bool, detail: str = "") -> None:
        ok = bool(condition)
        rows.append((label, "PASS" if ok else "FAIL",
                     [] if ok else [detail], {"Test": label, "Notes": detail or ""}))

    # Path A — new code
    db = _fresh_db()
    auto_learn_from_diagnosis(db, fault, diagnostic)
    doc = db["knowledge_base"].find_one({"code": "SPN 8888"})
    _case("Path A: document created",     doc is not None)
    _case("Path A: source=auto_learned",  doc and doc.get("source") == "auto_learned")
    _case("Path A: meaning set",          doc and doc.get("meaning") == "Gear ratio error")
    _case("Path A: severity=High",        doc and doc.get("severity") == "High")

    # Path A — existing auto_learned NOT overwritten
    auto_learn_from_diagnosis(db, fault, {**diagnostic, "purpose": "SHOULD NOT OVERWRITE"})
    doc = db["knowledge_base"].find_one({"code": "SPN 8888"})
    _case("Path A: existing auto_learned not overwritten",
          doc and doc.get("meaning") == "Gear ratio error",
          f"meaning was {doc.get('meaning') if doc else None}")

    # Path B — upgrade extracted_from_doc in-place
    db = _fresh_db()
    _seed(db, "SPN 8888", source="extracted_from_doc", meaning="", severity="Low")
    auto_learn_from_diagnosis(db, fault, diagnostic)
    doc = db["knowledge_base"].find_one({"code": "SPN 8888"})
    _case("Path B: source promoted to auto_learned", doc and doc.get("source") == "auto_learned")
    _case("Path B: meaning upgraded",               doc and doc.get("meaning") == "Gear ratio error")

    # Seed entry never touched
    db = _fresh_db()
    _seed(db, "SPN 8888", meaning="Original seed meaning")
    auto_learn_from_diagnosis(db, fault, diagnostic)
    doc = db["knowledge_base"].find_one({"code": "SPN 8888"})
    _case("seed entry meaning not overwritten",
          doc and doc.get("meaning") == "Original seed meaning",
          f"meaning was {doc.get('meaning') if doc else None}")

    return rows


# ============================================================================
# Console print
# ============================================================================

def _print_rows(section: str, rows: list[tuple]) -> int:
    print(f"=== {section} ===")
    passed = 0
    for label, status, diffs, _ in rows:
        print(f"  [{status}] {label}")
        for d in diffs:
            print(f"         x {d}")
        if status == "PASS":
            passed += 1
    return passed


# ============================================================================
# Doc generation
# ============================================================================

def _add_section(doc: Document, title: str, rows: list[tuple]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Test Case", "Status", "Notes"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for label, status, diffs, fields in rows:
        row = table.add_row().cells
        row[0].text = fields.get("Test", label)
        status_run = row[1].paragraphs[0].add_run(status)
        status_run.bold = True
        status_run.font.color.rgb = _GREEN if status == "PASS" else _RED
        row[2].text = "; ".join(diffs) if diffs else ""

    doc.add_paragraph()


def generate_doc(
    lu_rows: list[tuple],
    inc_rows: list[tuple],
    ext_rows: list[tuple],
    aln_rows: list[tuple],
    passed: int,
    total: int,
) -> None:
    os.makedirs(os.path.dirname(DOC_PATH), exist_ok=True)
    doc = Document()

    title = doc.add_heading("Knowledge Base — Test Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Module: core/knowledge_base.py  |  DB: mongomock (in-memory)")
    p = doc.add_paragraph()
    run = p.add_run(f"Result: {passed}/{total} passed")
    run.bold = True
    run.font.color.rgb = _GREEN if passed == total else _RED
    doc.add_paragraph()

    _add_section(doc, "lookup",                           lu_rows)
    _add_section(doc, "increment_occurrence",             inc_rows)
    _add_section(doc, "extract_and_insert_from_document", ext_rows)
    _add_section(doc, "auto_learn_from_diagnosis",        aln_rows)

    doc.save(DOC_PATH)
    print(f"Report written to {DOC_PATH}")


# ============================================================================
# Main
# ============================================================================

def main() -> None:
    lu_rows  = test_lookup()
    inc_rows = test_increment_occurrence()
    ext_rows = test_extract_and_insert()
    aln_rows = test_auto_learn()

    lu_pass  = _print_rows("lookup",                           lu_rows)
    inc_pass = _print_rows("increment_occurrence",             inc_rows)
    ext_pass = _print_rows("extract_and_insert_from_document", ext_rows)
    aln_pass = _print_rows("auto_learn_from_diagnosis",        aln_rows)

    total_pass = lu_pass + inc_pass + ext_pass + aln_pass
    total      = len(lu_rows) + len(inc_rows) + len(ext_rows) + len(aln_rows)
    print(f"\n{total_pass}/{total} passed")

    generate_doc(lu_rows, inc_rows, ext_rows, aln_rows, total_pass, total)


if __name__ == "__main__":
    main()

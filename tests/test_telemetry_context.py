"""
Telemetry Context — unit test suite.

Tests build_telemetry_snapshot and adjust_severity with fixed input/output pairs.
Writes results to tests/reports/report_telemetry_context.docx.

Run from diagnostic_agent/:
    conda run -n driverbook python tests/test_telemetry_context.py
"""

import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.telemetry_context import build_telemetry_snapshot, adjust_severity

DOC_PATH = os.path.join(os.path.dirname(__file__), "reports", "report_telemetry_context.docx")

_SENTINEL = -6.128e18
_GREEN = RGBColor(0x00, 0xAA, 0x55)
_RED   = RGBColor(0xCC, 0x00, 0x00)

# ============================================================================
# Test cases
# ============================================================================

SNAPSHOT_CASES: list[tuple] = [
    (
        {"engineCoolantTemperature": 90.0, "engineOilPressure": 35.0, "speed": 60.0},
        {"engineCoolantTemperature": 90.0, "engineOilPressure": 35.0, "speed": 60.0},
    ),
    (
        {"engineCoolantTemperature": _SENTINEL, "speed": 55.0},
        {"engineCoolantTemperature": None, "speed": 55.0},
    ),
    (
        {"defLevel": "2.5", "fuelLevel": "75"},
        {"defLevel": 2.5, "fuelLevel": 75.0},
    ),
    (
        {"engineOilPressure": "not_a_number"},
        {"engineOilPressure": None},
    ),
    (
        {},
        {"engineCoolantTemperature": None, "engineOilPressure": None,
         "speed": None, "fuelLevel": None, "defLevel": None, "engineSpeed": None},
    ),
]

SEVERITY_CASES: list[tuple] = [
    # base, fault, telemetry, expected, rule_note
    ("Low",    {"code": "SPN 110",  "ecu": "Engine Controller"}, {"engineCoolantTemperature": 110.0}, "Medium",   "Rule 1: hot coolant + engine"),
    ("Medium", {"code": "SPN 110",  "ecu": "Engine Controller"}, {"engineCoolantTemperature": 110.0}, "High",     "Rule 1: escalates from Medium"),
    ("High",   {"code": "SPN 110",  "ecu": "Engine Controller"}, {"engineCoolantTemperature": 110.0}, "Critical", "Rule 1: caps at Critical"),
    ("Low",    {"code": "SPN 110",  "ecu": "Engine Controller"}, {"engineCoolantTemperature": 95.0},  "Low",      "Rule 1: below threshold, no change"),
    ("Low",    {"code": "SPN 110",  "ecu": "Body Controller"},   {"engineCoolantTemperature": 110.0}, "Low",      "Rule 1: non-engine ECU, no change"),
    ("Low",    {"code": "SPN 100",  "ecu": "Engine #1"},         {"engineOilPressure": 15.0},         "Critical", "Rule 2: low oil + engine"),
    ("High",   {"code": "SPN 100",  "ecu": "Engine #1"},         {"engineOilPressure": 15.0},         "Critical", "Rule 2: always Critical regardless of base"),
    ("Low",    {"code": "SPN 1761", "ecu": "Aftertreatment"},    {"defLevel": 3.0},                   "Medium",   "Rule 3: low DEF + DEF code"),
    ("Medium", {"code": "SPN 4374", "ecu": "Emission Control"},  {"defLevel": 2.0},                   "High",     "Rule 3: escalates from Medium"),
    ("Low",    {"code": "SPN 521",  "ecu": "Aftertreatment"},    {"defLevel": 2.0},                   "Low",      "Rule 3: low DEF but non-DEF code"),
    ("Unknown",{"code": "SPN 110",  "ecu": "Engine Controller"}, {"engineCoolantTemperature": 110.0}, "Medium",   "Unknown base defaults to Low first"),
    ("High",   {"code": "SPN 110",  "ecu": "Engine Controller"}, {},                                  "High",     "No telemetry, no change"),
]


# ============================================================================
# Runners — return list of (label, status, diffs, fields)
# ============================================================================

def run_snapshot_cases() -> list[tuple]:
    rows = []
    for raw, expected_subset in SNAPSHOT_CASES:
        result = build_telemetry_snapshot(raw)
        diffs = [
            f"{k}: expected={ev!r} actual={result.get(k)!r}"
            for k, ev in expected_subset.items()
            if result.get(k) != ev
        ]
        ok = len(diffs) == 0
        keys_str = ", ".join(expected_subset.keys())
        rows.append((
            f"keys=[{keys_str}]",
            "PASS" if ok else "FAIL",
            diffs,
            {"Input keys": str(list(raw.keys())), "Checked keys": keys_str},
        ))
    return rows


def run_severity_cases() -> list[tuple]:
    rows = []
    for base, fault, telemetry, expected, note in SEVERITY_CASES:
        result = adjust_severity(base, fault, telemetry)
        ok = result == expected
        label = f"{note}"
        diffs = [] if ok else [f"expected={expected!r}  actual={result!r}"]
        rows.append((
            label,
            "PASS" if ok else "FAIL",
            diffs,
            {
                "Rule / Note": note,
                "Base": base,
                "ECU": fault.get("ecu", ""),
                "Code": fault.get("code", ""),
                "Expected": expected,
                "Actual": result,
            },
        ))
    return rows


# ============================================================================
# Console print
# ============================================================================

def _print_rows(section: str, rows: list[tuple]) -> int:
    print(f"=== {section} ===")
    passed = 0
    for label, status, diffs, _ in rows:
        print(f"  [{status}] {label}")
        for d in diffs:
            print(f"         x {d}")
        if status == "PASS":
            passed += 1
    return passed


# ============================================================================
# Doc generation
# ============================================================================

def _add_section(doc: Document, title: str, rows: list[tuple], col_keys: list[str]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(col_keys) + 2)
    table.style = "Table Grid"
    headers = col_keys + ["Status", "Notes"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for label, status, diffs, fields in rows:
        row = table.add_row().cells
        for i, key in enumerate(col_keys):
            row[i].text = str(fields.get(key, ""))
        status_cell = row[len(col_keys)]
        run = status_cell.paragraphs[0].add_run(status)
        run.bold = True
        run.font.color.rgb = _GREEN if status == "PASS" else _RED
        row[len(col_keys) + 1].text = "; ".join(diffs) if diffs else ""

    doc.add_paragraph()


def generate_doc(snap_rows: list[tuple], sev_rows: list[tuple], passed: int, total: int) -> None:
    os.makedirs(os.path.dirname(DOC_PATH), exist_ok=True)
    doc = Document()

    title = doc.add_heading("Telemetry Context — Test Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Module: core/telemetry_context.py")
    p = doc.add_paragraph()
    run = p.add_run(f"Result: {passed}/{total} passed")
    run.bold = True
    run.font.color.rgb = _GREEN if passed == total else _RED
    doc.add_paragraph()

    _add_section(doc, "build_telemetry_snapshot", snap_rows, ["Input keys", "Checked keys"])
    _add_section(doc, "adjust_severity", sev_rows, ["Rule / Note", "Base", "ECU", "Code", "Expected", "Actual"])

    doc.save(DOC_PATH)
    print(f"Report written to {DOC_PATH}")


# ============================================================================
# Main
# ============================================================================

def main() -> None:
    snap_rows = run_snapshot_cases()
    sev_rows  = run_severity_cases()

    snap_pass = _print_rows("build_telemetry_snapshot", snap_rows)
    sev_pass  = _print_rows("adjust_severity",          sev_rows)

    total_pass = snap_pass + sev_pass
    total      = len(snap_rows) + len(sev_rows)
    print(f"\n{total_pass}/{total} passed")

    generate_doc(snap_rows, sev_rows, total_pass, total)


if __name__ == "__main__":
    main()
